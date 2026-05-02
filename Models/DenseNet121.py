# ==========================================
# Google Colab version of densenet121.py
# ==========================================

# Install required library (not preinstalled in Colab)
!pip install python-docx

# Imports
import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import DenseNet121
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout, BatchNormalization
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ModelCheckpoint, ReduceLROnPlateau
from sklearn.metrics import classification_report, confusion_matrix
import seaborn as sns
from docx import Document
from docx.shared import Inches

# ==========================================
# Mount Google Drive
# ==========================================

# Paths (update these according to your dataset in Drive)
input_path = "/content/drive/MyDrive/Final Dataset"   # dataset folder
output_path = "/content/drive/MyDrive/Aquarium Fish/Densenet"
os.makedirs(output_path, exist_ok=True)

# Parameters
EPOCHS = 25
BATCH_SIZE = 32
LEARNING_RATE = 0.0001

# Create Document
doc = Document()
doc.add_heading('Classification Report for DenseNet121', 0)

# ==========================================
# DenseNet121 model builder
# ==========================================
def create_densenet_model(input_shape, num_classes):
    base_model = DenseNet121(weights='imagenet', include_top=False, input_shape=input_shape)

    # Fine-tuning last 50 layers
    for layer in base_model.layers[:-50]:
        layer.trainable = False
    for layer in base_model.layers[-50:]:
        layer.trainable = True

    x = base_model.output
    x = GlobalAveragePooling2D()(x)
    x = Dense(2048, activation='relu')(x)
    x = BatchNormalization()(x)
    x = Dropout(0.5)(x)
    predictions = Dense(num_classes, activation='softmax')(x)

    model = Model(inputs=base_model.input, outputs=predictions)
    return model

# ==========================================
# Data pipeline
# ==========================================
def process_data(input_path, img_size, batch_size):
    datagen = ImageDataGenerator(
        preprocessing_function=tf.keras.applications.densenet.preprocess_input,
        validation_split=0.2,
        rotation_range=30,
        width_shift_range=0.2,
        height_shift_range=0.2,
        shear_range=0.2,
        zoom_range=0.2,
        horizontal_flip=True,
        fill_mode="nearest",
        brightness_range=[0.8, 1.2],
        channel_shift_range=20.0
    )

    train_gen = datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',
        subset='training',
        shuffle=True
    )

    val_gen = datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',
        subset='validation',
        shuffle=False
    )

    return train_gen, val_gen

# Image size for DenseNet121
img_size = (299, 299)

# Prepare data
train_gen, val_gen = process_data(input_path, img_size, BATCH_SIZE)

# Create model
model = create_densenet_model((img_size[0], img_size[1], 3), train_gen.num_classes)

# Compile
model.compile(optimizer=Adam(learning_rate=LEARNING_RATE),
              loss='categorical_crossentropy',
              metrics=['accuracy'])

# Callbacks
checkpoint = ModelCheckpoint(os.path.join(output_path, 'DenseNet121_best_model.keras'),
                             monitor='val_accuracy',
                             save_best_only=True,
                             mode='max',
                             verbose=1)

reduce_lr = ReduceLROnPlateau(monitor='val_loss', factor=0.1, patience=3,
                              verbose=1, min_lr=1e-6)

# ==========================================
# Training
# ==========================================
history = model.fit(
    train_gen,
    epochs=EPOCHS,
    alidation_data=val_gen,
    callbacks=[checkpoint, reduce_lr]
)

# ==========================================
# Accuracy & Loss Plots
# ==========================================
plt.figure(figsize=(12, 5))

plt.subplot(1, 2, 1)
plt.plot(history.history['accuracy'], label='Train Accuracy', marker='o')
plt.plot(history.history['val_accuracy'], label='Validation Accuracy', marker='o')
plt.title("DenseNet121 Accuracy")
plt.xlabel("Epochs")
plt.ylabel("Accuracy")
plt.legend()

plt.subplot(1, 2, 2)
plt.plot(history.history['loss'], label='Train Loss', marker='o')
plt.plot(history.history['val_loss'], label='Validation Loss', marker='o')
plt.title("DenseNet121 Loss")
plt.xlabel("Epochs")
plt.ylabel("Loss")
plt.legend()

acc_loss_path = os.path.join(output_path, 'DenseNet121_accuracy_loss_curve.png')
plt.savefig(acc_loss_path)
plt.close()

# ==========================================
# Evaluation & Predictions
# ==========================================
model.load_weights(os.path.join(output_path, 'DenseNet121_best_model.keras'))
eval_result = model.evaluate(val_gen)

Y_pred = model.predict(val_gen)
y_pred = np.argmax(Y_pred, axis=1)
y_true = val_gen.classes

report = classification_report(y_true, y_pred, target_names=list(train_gen.class_indices.keys()))

# Confusion Matrix
cm = confusion_matrix(y_true, y_pred)
cm_path = os.path.join(output_path, 'DenseNet121_confusion_matrix.png')

plt.figure(figsize=(12, 10))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=train_gen.class_indices.keys(),
            yticklabels=train_gen.class_indices.keys(),
            annot_kws={"size": 14})
plt.xticks(rotation=45, fontsize=12)
plt.yticks(rotation=0, fontsize=12)
plt.xlabel('Predicted', fontsize=14)
plt.ylabel('Actual', fontsize=14)
plt.title('DenseNet121 Confusion Matrix', fontsize=16)
plt.tight_layout()
plt.savefig(cm_path)
plt.close()

# ==========================================
# Word Report
# ==========================================
doc.add_heading('DenseNet121 Performance', level=1)
doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

doc.add_heading('Classification Report', level=2)
doc.add_paragraph(report)

doc.add_heading('Training Accuracy and Loss', level=2)
doc.add_picture(acc_loss_path, width=Inches(6))

doc.add_heading('Confusion Matrix', level=2)
doc.add_picture(cm_path, width=Inches(6))

doc.save(os.path.join(output_path, 'DenseNet121_classification_report.docx'))

print("✅ DenseNet121 model training and report generation completed! Results saved in Google Drive.")